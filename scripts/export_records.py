#!/usr/bin/env python3
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, List, Dict, Optional


_PRIMARY_RECORD_KEYS = (
    "entries",
    "items",
    "line_items",
    "records",
    "rows",
    "results",
    "data",
)


def _flatten_mapping(data: Dict[str, Any], *, prefix: str = "") -> Dict[str, Any]:
    """Flatten nested dicts into dotted keys and stringify complex values."""
    flat: Dict[str, Any] = {}
    for key, value in data.items():
        new_key = f"{prefix}.{key}" if prefix else key
        if isinstance(value, dict):
            flat.update(_flatten_mapping(value, prefix=new_key))
        elif isinstance(value, list):
            if not value:
                flat[new_key] = ""
            elif all(isinstance(item, (str, int, float, bool)) or item is None for item in value):
                flat[new_key] = ", ".join("" if item is None else str(item) for item in value)
            else:
                flat[new_key] = json.dumps(value, ensure_ascii=False)
        else:
            flat[new_key] = value
    return flat


def _select_primary_records(payload: Dict[str, Any]) -> Optional[str]:
    """Pick a list-of-dict field to explode into rows, if present."""
    candidates: List[str] = []
    for key, value in payload.items():
        if isinstance(value, list):
            if value and all(isinstance(item, dict) for item in value):
                candidates.append(key)
            elif not value and key in _PRIMARY_RECORD_KEYS:
                # Empty list for a preferred key – still treat as candidate
                candidates.append(key)
    if not candidates:
        return None
    for preferred in _PRIMARY_RECORD_KEYS:
        if preferred in candidates:
            return preferred
    return candidates[0]


def ensure_records(obj: Any) -> List[Dict[str, Any]]:
    """Normalize arbitrary data into flat row dictionaries for export.

    - list[dict]: flatten each mapping individually
    - list[scalar/...]: wrap values under "value"
    - dict: flatten, exploding a primary list-of-dict field (entries/items/etc.) into rows
    - str: wrap under "text"
    - other: wrap under "value"
    """
    if isinstance(obj, list):
        if all(isinstance(item, dict) for item in obj):
            return [_flatten_mapping(item) for item in obj]
        return [{"value": item} for item in obj]

    if isinstance(obj, dict):
        primary_key = _select_primary_records(obj)

        if primary_key:
            base = {key: value for key, value in obj.items() if key != primary_key}
            base_flat = _flatten_mapping(base)
            rows: List[Dict[str, Any]] = []
            primary_list = obj.get(primary_key) or []

            if isinstance(primary_list, list) and primary_list:
                for entry in primary_list:
                    if isinstance(entry, dict):
                        row = {**base_flat}
                        row.update(_flatten_mapping(entry, prefix=primary_key))
                        rows.append(row)
                    else:
                        row = {**base_flat, primary_key: entry}
                        rows.append(row)
                if rows:
                    return rows
            # Fall back to a single row even if the primary list is empty
            return [base_flat]

        return [_flatten_mapping(obj)]

    if isinstance(obj, str):
        return [{"text": obj}]

    return [{"value": obj}]


def all_columns(records: Iterable[Dict[str, Any]]) -> List[str]:
    cols: List[str] = []
    for r in records:
        for k in r.keys():
            if k not in cols:
                cols.append(k)
    return cols


def _norm_cell(v: Any) -> Any:
    if v is None:
        return ""
    if isinstance(v, (str, int, float, bool)):
        return v
    # JSON-encode nested structures
    return json.dumps(v, ensure_ascii=False)


def to_json_bytes(records: List[Dict[str, Any]]) -> bytes:
    return json.dumps(records, ensure_ascii=False, indent=2).encode("utf-8")


def to_markdown_bytes(records: List[Dict[str, Any]], cols: List[str]) -> bytes:
    def esc(s: str) -> str:
        return s.replace("|", "\\|").replace("\n", "<br>")

    lines: List[str] = []
    lines.append("| " + " | ".join(cols) + " |")
    lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
    for r in records:
        row = [esc(str(_norm_cell(r.get(c, "")))) for c in cols]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines).encode("utf-8")


def to_xlsx_bytes(records: List[Dict[str, Any]], cols: List[str]) -> bytes:
    from openpyxl import Workbook  # requires dependency

    wb = Workbook()
    ws = wb.active
    ws.title = "Export"
    ws.append(cols)
    for r in records:
        ws.append([_norm_cell(r.get(c, "")) for c in cols])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_docx_bytes(records: List[Dict[str, Any]], cols: List[str]) -> bytes:
    """Create a simple DOCX table for records.

    - First row as headers
    - One row per record; non-primitive values are JSON-encoded
    """
    from docx import Document  # requires python-docx

    doc = Document()
    doc.add_heading("Export", level=1)

    # Create table with header
    table = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr_cells[i].text = str(c)

    for r in records:
        row_cells = table.add_row().cells
        for i, c in enumerate(cols):
            row_cells[i].text = str(_norm_cell(r.get(c, "")))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_docx_from_text_bytes(text: str, title: str | None = None) -> bytes:
    """Create a DOCX document from plain text (best-effort for Markdown)."""
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    # Keep simple: one paragraph block; avoid complex markdown parsing
    for line in str(text or "").splitlines():
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _cli() -> None:
    import argparse

    p = argparse.ArgumentParser(description="Export records (JSON) to JSON/Markdown/XLSX")
    p.add_argument("input", help="Path to JSON file (array or object)")
    p.add_argument("--out", default="export", help="Output base name without extension")
    args = p.parse_args()

    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    recs = ensure_records(data)
    cols = all_columns(recs)

    Path(args.out + ".json").write_bytes(to_json_bytes(recs))
    Path(args.out + ".md").write_bytes(to_markdown_bytes(recs, cols))
    Path(args.out + ".xlsx").write_bytes(to_xlsx_bytes(recs, cols))
    print(f"Wrote {args.out}.json {args.out}.md {args.out}.xlsx")


if __name__ == "__main__":
    _cli()
